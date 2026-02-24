import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor


TIMESTAMP_LINE_RE = re.compile(
    r"^\[(\d{1,2}:\d{2}:\d{2})\]\s+(Speaker\s+\w+):\s*$"
)


def parse_transcript(txt_path: str):
    """
    Parse the .txt transcript produced by transcribe.py into
    a list of (timestamp, speaker, text) utterances.
    """
    with open(txt_path, "r", encoding="utf-8") as f:
        lines = [line.rstrip("\n") for line in f]

    entries = []
    i = 0

    # Skip header lines until we hit the first timestamp line
    while i < len(lines) and not TIMESTAMP_LINE_RE.match(lines[i]):
        i += 1

    while i < len(lines):
        match = TIMESTAMP_LINE_RE.match(lines[i])
        if not match:
            i += 1
            continue

        timestamp = match.group(1)
        speaker = match.group(2)
        i += 1

        utter_lines = []
        # Collect non-empty lines until next timestamp or blank separator
        while (
            i < len(lines)
            and lines[i].strip() != ""
            and not TIMESTAMP_LINE_RE.match(lines[i])
        ):
            utter_lines.append(lines[i].strip())
            i += 1

        text = " ".join(utter_lines).strip()
        entries.append((timestamp, speaker, text))

        # Skip any blank separator lines
        while i < len(lines) and lines[i].strip() == "":
            i += 1

    return entries


def configure_document_styles(document: Document):
    """
    Configure basic page layout and fonts for a clean, readable therapy transcript.
    """
    # Page layout: standard 1-inch margins
    for section in document.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Base font
    normal_style = document.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Calibri"
    normal_font.size = Pt(11)


def create_docx_from_transcript(txt_path: str, docx_path: str):
    entries = parse_transcript(txt_path)

    if not entries:
        raise ValueError("No transcript entries were parsed from the .txt file.")

    document = Document()
    configure_document_styles(document)

    base_name = os.path.basename(txt_path)
    session_title = "Therapy Session Transcript"

    # Title
    title = document.add_paragraph(session_title)
    title.style = document.styles["Title"]

    # Optional metadata / subtitle
    subtitle_text = f"Source file: {base_name}"
    try:
        subtitle = document.add_paragraph(subtitle_text)
        subtitle.style = document.styles["Subtitle"]
    except KeyError:
        # Fallback if the "Subtitle" style is not available
        subtitle = document.add_paragraph(subtitle_text)
        subtitle.bold = True

    # Session date (derived from file last modified time, if available)
    try:
        mtime = datetime.fromtimestamp(os.path.getmtime(txt_path))
        meta = document.add_paragraph()
        meta.add_run("Session date (file timestamp): ").bold = True
        meta.add_run(mtime.strftime("%Y-%m-%d %H:%M"))
    except OSError:
        pass

    document.add_paragraph()

    # Introductory note
    note = document.add_paragraph()
    note_run = note.add_run(
        "Note: This transcript is anonymized and intended for clinical/educational use. "
        "Speakers are labeled generically (e.g., Speaker A, Speaker B)."
    )
    note_run.italic = True

    document.add_paragraph()  # Blank line for spacing

    # Color palette to differentiate speakers.
    # Will be reused/cycled automatically if there are many speakers.
    color_palette = [
        RGBColor(31, 73, 125),   # dark blue
        RGBColor(79, 129, 189),  # blue
        RGBColor(112, 48, 160),  # purple
        RGBColor(0, 128, 0),     # green
        RGBColor(192, 80, 77),   # red
        RGBColor(128, 96, 0),    # brown/gold
    ]
    speaker_colors = {}

    # Add each utterance as a timestamp + speaker heading, followed by the text
    for timestamp, speaker, text in entries:
        # Assign a color for this speaker if not already assigned
        if speaker not in speaker_colors:
            color = color_palette[len(speaker_colors) % len(color_palette)]
            speaker_colors[speaker] = color
        else:
            color = speaker_colors[speaker]

        # Heading line with timestamp and speaker
        heading_para = document.add_paragraph()
        heading_run = heading_para.add_run(f"[{timestamp}] {speaker}")
        heading_run.bold = True
        heading_run.font.color.rgb = color

        # Utterance text (same color as the speaker heading)
        if text:
            text_para = document.add_paragraph()
            text_run = text_para.add_run(text)
            text_para.style = document.styles["Normal"]
            text_run.font.color.rgb = color

        document.add_paragraph()  # Blank line between turns

    document.save(docx_path)


def main():
    if len(sys.argv) < 2:
        print(
            "Usage: python txt_to_docx.py path\\to\\transcript.txt [output.docx]",
            file=sys.stderr,
        )
        sys.exit(1)

    txt_path = sys.argv[1]

    if not os.path.exists(txt_path):
        print(f"Error: file not found: {txt_path}", file=sys.stderr)
        sys.exit(1)

    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
    else:
        base, _ = os.path.splitext(txt_path)
        docx_path = base + ".docx"

    try:
        create_docx_from_transcript(txt_path, docx_path)
    except Exception as exc:
        print(f"Failed to create .docx: {exc}", file=sys.stderr)
        sys.exit(1)

    print(f"Created Word document: {docx_path}")


if __name__ == "__main__":
    main()

